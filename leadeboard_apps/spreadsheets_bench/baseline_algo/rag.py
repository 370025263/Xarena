# rag.py — Milvus-Lite 版本：本地嵌入式向量库 + 朴素 RAG（pypdf 分页、按 token 分块）
# 说明：
# - 维持你之前的 LLM 网络/重试与 Tokenizer 初始化风格；
# - 嵌入仍走 OpenAI 兼容 /v1/embeddings（EMBEDDING_API_ENDPOINT / EMBEDDING_MODEL / EMBEDDING_API_KEY）；
# - 检索存储使用 Milvus-Lite（pymilvus 的 MilvusClient，uri 指向本地路径即可）。

from __future__ import annotations
import os
import io
import re
import math
import glob
import json
import time
import atexit
import asyncio
import random
import threading
from dataclasses import dataclass, field
from typing import Any, Dict, List, Optional, Tuple

import requests
import httpx
import numpy as np
from openai import AsyncOpenAI
from openai import APIConnectionError, APIStatusError, RateLimitError, APIError

# ---------- Milvus ----------
from pymilvus import MilvusClient, DataType

# ---------- 文档解析 ----------
# 1. PDF 解析
try:
    from pypdf import PdfReader
except Exception:
    PdfReader = None  # 若未安装，会在 build_index 阶段给出明确报错

# 2. DOCX 解析 (新增)
try:
    import docx
except Exception:
    docx = None

# ---------- Tokenizer：遵循你给的初始化约定 ----------
# 你的约定：
#   if self.tokenizer is None:
#       if self.tiktoken_model_name:
#           self.tokenizer = TiktokenTokenizer(self.tiktoken_model_name)
#       else:
#           self.tokenizer = TiktokenTokenizer()
#   tiktoken_model_name default = "gpt-4o-mini"
class TiktokenTokenizer:
    def __init__(self, model_name: Optional[str] = None):
        try:
            import tiktoken as _t
        except Exception as e:
            raise RuntimeError("tiktoken 未安装，请在镜像中安装 tiktoken") from e

        # 优先用 encoding_for_model，失败时回落到常见 base
        self._t = _t
        enc = None
        if model_name:
            try:
                enc = _t.encoding_for_model(model_name)
            except Exception:
                pass
        if enc is None:
            # gpt-4o 系列一般可用 o200k_base；不行再退 cl100k_base
            try:
                enc = _t.get_encoding("o200k_base")
            except Exception:
                enc = _t.get_encoding("cl100k_base")
        self._enc = enc

    def count(self, s: str) -> int:
        return len(self._enc.encode(s or ""))

    def truncate(self, s: str, max_tokens: int) -> str:
        ids = self._enc.encode(s or "")
        if len(ids) <= max_tokens:
            return s
        return self._enc.decode(ids[:max_tokens])


# ==============================
# 工具函数
# ==============================
def _l2_normalize(v: np.ndarray, eps: float = 1e-12) -> np.ndarray:
    nrm = np.linalg.norm(v, axis=-1, keepdims=True)
    return v / np.clip(nrm, eps, None)


def _read_file_text(path: str) -> List[Tuple[str, Dict[str, Any]]]:
    """
    读取文件并返回 [(文本, 元数据), ...]
    - PDF：按页返回
    - DOCX：整文件拼接 (page=-1)
    - TXT/MD：整文件一条
    """
    ext = os.path.splitext(path)[1].lower()
    out: List[Tuple[str, Dict[str, Any]]] = []

    if ext == ".pdf":
        if PdfReader is None:
            raise RuntimeError("未安装 pypdf，请安装后再解析 PDF")
        try:
            r = PdfReader(path)
            for i, page in enumerate(r.pages):
                try:
                    txt = page.extract_text() or ""
                except Exception:
                    txt = ""
                md = {"path": path, "doc_id": os.path.basename(path), "page": i}
                if txt.strip():
                    out.append((txt, md))
        except Exception as e:
            raise RuntimeError(f"解析 PDF 失败: {path} : {e}")

    elif ext == ".docx":
        if docx is None:
            raise RuntimeError("未安装 python-docx，请安装后再解析 DOCX")
        try:
            doc = docx.Document(path)
            full_text = []
            # 提取段落
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            # 简单提取表格（可选， naive 模式下简单拼接）
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells if cell.text.strip()]
                    if row_text:
                        full_text.append(" | ".join(row_text))

            txt = "\n\n".join(full_text)
            if txt.strip():
                # DOCX 没有严格固定分页，统一 page=-1
                out.append((txt, {"path": path, "doc_id": os.path.basename(path), "page": -1}))
        except Exception as e:
            raise RuntimeError(f"解析 DOCX 失败: {path} : {e}")

    elif ext in (".txt", ".md"):
        try:
            with io.open(path, "r", encoding="utf-8", errors="ignore") as f:
                txt = f.read()
            if txt.strip():
                out.append((txt, {"path": path, "doc_id": os.path.basename(path), "page": -1}))
        except Exception as e:
            raise RuntimeError(f"读取文本失败: {path} : {e}")
    else:
        # 其他格式此处略过
        pass

    return out


def _sliding_token_chunks(
    text: str,
    tokenizer: TiktokenTokenizer,
    chunk_tokens: int,
    overlap_tokens: int,
) -> List[str]:
    """按 token 滑窗切分文本"""
    ids = tokenizer._enc.encode(text or "")
    n = len(ids)
    if n == 0:
        return []
    chunks = []
    step = max(1, chunk_tokens - overlap_tokens)
    for start in range(0, n, step):
        end = start + chunk_tokens
        sub = ids[start:end]
        if not sub:
            break
        chunks.append(tokenizer._enc.decode(sub))
        if end >= n:
            break
    return chunks


# ==============================
# Engine
# ==============================
def load_engine(config: Dict[str, Any]) -> "Engine":
    return Engine(config)


@dataclass
class Engine:
    # ---- 基础配置（与 app.py 保持一致的键名）----
    cfg: Dict[str, Any]
    tokenizer: Optional[TiktokenTokenizer] = None
    tiktoken_model_name: str = field(default="gpt-4o-mini")

    # ---- LLM 重试参数（沿用你之前逻辑的名字）----
    _max_retries: int = field(default=int(os.getenv("LLM_MAX_RETRIES", "3")))
    _backoff_base: float = field(default=float(os.getenv("LLM_RETRY_BACKOFF_BASE", "0.6")))
    _backoff_cap: float = field(default=float(os.getenv("LLM_RETRY_BACKOFF_CAP", "6")))
    _retry_status: set = field(default_factory=lambda: {
        int(s) for s in os.getenv("LLM_RETRY_STATUS", "429,500,502,503,504").split(",") if s.strip().isdigit()
    })
    _req_timeout: float = field(default=float(os.getenv("LLM_REQ_TIMEOUT", "60")))
    _sdk_debug: bool = field(default=os.getenv("LLM_SDK_DEBUG", "0") == "1")

    # ---- RAG 参数（环境变量可配）----
    chunk_tokens: int = field(default=int(os.getenv("RAG_CHUNK_TOKENS", "512")))
    chunk_overlap: int = field(default=int(os.getenv("RAG_CHUNK_OVERLAP", "50")))
    top_k: int = field(default=int(os.getenv("RAG_TOPK", "4")))
    context_max_tokens: int = field(default=int(os.getenv("RAG_CONTEXT_MAX_TOKENS", "1200")))
    text_max_chars: int = field(default=int(os.getenv("RAG_TEXT_MAX_CHARS", "8192")))
    embed_batch: int = field(default=int(os.getenv("RAG_EMBED_BATCH", "16")))
    metric_type: str = field(default=os.getenv("RAG_METRIC", "COSINE").upper())  # COSINE / IP / L2
    index_type: str = field(default=os.getenv("RAG_INDEX_TYPE", "AUTOINDEX").upper())  # AUTOINDEX/IVF_FLAT/HNSW/FLAT
    score_threshold: float = field(default=float(os.getenv("RAG_MIN_SCORE", "0.0")))

    # ---- Milvus Lite 参数 ----
    milvus_uri: str = field(default=os.getenv("MILVUS_URI", os.getenv("RAG_DB_PATH", "/app/rag_storage/milvus.db")))
    collection: str = field(default=os.getenv("MILVUS_COLLECTION", "kb_default"))

    # ---- 内部状态 ----
    _loop: Optional[asyncio.AbstractEventLoop] = field(default=None, init=False)
    _loop_thread: Optional[threading.Thread] = field(default=None, init=False)
    _loop_started: bool = field(default=False, init=False)
    _loop_lock: threading.Lock = field(default_factory=threading.Lock, init=False)
    _llm_client: Optional[AsyncOpenAI] = field(default=None, init=False)
    _emb_sess: Optional[requests.Session] = field(default=None, init=False)
    _last_usage: Dict[str, int] = field(default_factory=lambda: {"prompt_tokens": 0, "completion_tokens": 0, "total_tokens": 0})
    _stage: str = field(default="idle", init=False)

    # ---- Milvus Client ----
    _m: Optional[MilvusClient] = field(default=None, init=False)
    _dim: int = field(default=0, init=False)

    def __post_init__(self):
        # Tokenizer 初始化（遵循你的约定）
        if self.tokenizer is None:
            if self.tiktoken_model_name:
                self.tokenizer = TiktokenTokenizer(self.tiktoken_model_name)
            else:
                self.tokenizer = TiktokenTokenizer()

        # 维持 tiktoken 缓存位置
        os.environ.setdefault("TIKTOKEN_CACHE_DIR", "/models/tiktoken/encoders/")

        # 取 embedding 维度
        self._dim = int(self.cfg.get("embedding_dim", 1024))

        atexit.register(self.close)
        print(f"[Engine] Milvus-Lite RAG, collection={self.collection}, uri={self.milvus_uri}, dim={self._dim}, metric={self.metric_type}, index={self.index_type}")

    # ----------------- 事件循环 -----------------
    def _loop_worker(self) -> None:
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        self._loop = loop
        try:
            loop.run_forever()
        finally:
            pending = asyncio.all_tasks(loop)
            for t in pending:
                t.cancel()
            try:
                loop.run_until_complete(asyncio.gather(*pending, return_exceptions=True))
            except Exception:
                pass
            loop.close()

    def _start_loop_once(self) -> None:
        if self._loop_started and self._loop and self._loop_thread and self._loop_thread.is_alive():
            return
        with self._loop_lock:
            if self._loop_started and self._loop and self._loop_thread and self._loop_thread.is_alive():
                return
            self._loop_thread = threading.Thread(target=self._loop_worker, name="EngineEventLoop", daemon=True)
            self._loop_thread.start()
            while self._loop is None:
                pass
            self._loop_started = True

    def _run(self, coro):
        self._start_loop_once()
        assert self._loop is not None, "Event loop not initialized"
        fut = asyncio.run_coroutine_threadsafe(coro, self._loop)
        return fut.result()

    def close(self) -> None:
        if not self._loop_started or not self._loop:
            return
        try:
            self._loop.call_soon_threadsafe(self._loop.stop)
            if self._loop_thread and self._loop_thread.is_alive():
                self._loop_thread.join(timeout=3.0)
        except Exception:
            pass
        finally:
            self._loop_started = False
            self._loop = None
            self._loop_thread = None

    # ----------------- Milvus 准备 -----------------
    def _milvus(self) -> MilvusClient:
        if self._m is None:
            # uri 指向本地路径 => 自动启用 Milvus-Lite
            self._m = MilvusClient(uri=self.milvus_uri)
        return self._m

    def _ensure_collection(self) -> None:
        m = self._milvus()
        if not m.has_collection(self.collection):
            # 自定义 schema（需要额外字段：text/path/doc_id/page/chunk_id）
            schema = m.create_schema()
            schema.add_field(field_name="id", datatype=DataType.INT64, is_primary=True, auto_id=True)
            schema.add_field(field_name="vector", datatype=DataType.FLOAT_VECTOR, dim=self._dim)
            schema.add_field(field_name="text", datatype=DataType.VARCHAR, max_length=self.text_max_chars)
            schema.add_field(field_name="doc_id", datatype=DataType.VARCHAR, max_length=512)
            schema.add_field(field_name="path", datatype=DataType.VARCHAR, max_length=1024)
            schema.add_field(field_name="page", datatype=DataType.INT64)
            schema.add_field(field_name="chunk_id", datatype=DataType.INT64)

            index_params = m.prepare_index_params()
            # AUTOINDEX + metric（Milvus-Lite 友好）
            index_params.add_index(
                field_name="vector",
                index_type=self.index_type,
                metric_type=self.metric_type,
                params={},  # 交给 AUTOINDEX 或默认策略
            )
            m.create_collection(self.collection, schema=schema, index_params=index_params)
            print(f"[Engine] create Milvus collection: {self.collection}")
        else:
            # 已存在直接返回
            pass

    # ----------------- Embedding -----------------
    def _embed_batch(self, texts: List[str]) -> np.ndarray:
        """
        嵌入：OpenAI 兼容接口 /v1/embeddings
        - 端点：cfg['embedding_endpoint']
        - 模型：cfg['embedding_model']
        - 可选：EMBEDDING_API_KEY
        - 不继承系统代理（保持与之前“内网”调用一致）：Session.trust_env=False
        """
        if not texts:
            return np.zeros((0, self._dim), dtype=np.float32)

        if self._emb_sess is None:
            s = requests.Session()
            s.trust_env = False
            self._emb_sess = s

        url = (self.cfg.get("embedding_endpoint") or "").rstrip("/") + "/embeddings"
        model = self.cfg.get("embedding_model") or ""
        api_key = os.getenv("EMBEDDING_API_KEY", "")

        headers = {"Content-Type": "application/json"}
        if api_key:
            headers["Authorization"] = f"Bearer {api_key}"

        payload = {"model": model, "input": texts}
        r = self._emb_sess.post(url, headers=headers, json=payload, timeout=360)
        r.raise_for_status()
        data = r.json()

        vecs = [it["embedding"] for it in data.get("data", [])]
        arr = np.asarray(vecs, dtype=np.float32)
        if arr.shape[-1] != self._dim:
            raise RuntimeError(f"Embedding dim mismatch: got {arr.shape[-1]}, expected {self._dim}")
        if self.metric_type == "COSINE":
            arr = _l2_normalize(arr)
        return arr

    # ----------------- LLM -----------------
    async def _ensure_llm_async(self) -> None:
        if self._llm_client is None:
            # 关键修改：注入关闭 SSL 验证的 httpx client
            self._llm_client = AsyncOpenAI(
                base_url=self.cfg["llm_base_url"],
                api_key=self.cfg["llm_api_key"],
                http_client=httpx.AsyncClient(verify=False),
            )

    async def _llm_call(self, messages: List[Dict[str, str]], **kwargs) -> str:
        model = self.cfg["llm_model_name"]
        max_tokens = kwargs.get("max_tokens")
        temperature = kwargs.get("temperature")

        def _log_try(i: int, note: str) -> None:
            if self._sdk_debug:
                print(f"[llm][try {i+1}/{self._max_retries+1}] model={model} mt={max_tokens} temp={temperature} note={note}")

        last_exc: Optional[Exception] = None
        for i in range(self._max_retries + 1):
            _log_try(i, "sending")
            try:
                resp = await self._llm_client.chat.completions.create(  # type: ignore
                    model=model,
                    messages=messages,
                    max_tokens=max_tokens,
                    temperature=temperature,
                    stream=False,
                    extra_body={"enable_thinking": False},
                    timeout=self._req_timeout,
                )
                text = resp.choices[0].message.content or ""
                usage = getattr(resp, "usage", None)
                if usage is not None:
                    pt = int(getattr(usage, "prompt_tokens", 0) or 0)
                    ct = int(getattr(usage, "completion_tokens", 0) or 0)
                    self._last_usage = {
                        "prompt_tokens": pt,
                        "completion_tokens": ct,
                        "total_tokens": pt + ct,
                    }
                return text

            except APIStatusError as e:
                status = getattr(e, "status_code", None)
                if self._sdk_debug:
                    print(f"[llm][err] APIStatusError status={status} msg={getattr(e, 'message', e)}")
                last_exc = e
                if status in self._retry_status:
                    backoff = min(self._backoff_cap, self._backoff_base * (2**i))
                    backoff = backoff * (0.75 + 0.5 * random.random())
                    await asyncio.sleep(backoff)
                    continue
                break

            except (APIConnectionError, RateLimitError) as e:
                if self._sdk_debug:
                    print(f"[llm][err] {type(e).__name__}: {e}")
                last_exc = e
                backoff = min(self._backoff_cap, self._backoff_base * (2**i))
                backoff = backoff * (0.75 + 0.5 * random.random())
                await asyncio.sleep(backoff)
                continue

            except APIError as e:
                if self._sdk_debug:
                    print(f"[llm][err] APIError: {e}")
                last_exc = e
                backoff = min(self._backoff_cap, self._backoff_base * (2**i))
                backoff = backoff * (0.75 + 0.5 * random.random())
                await asyncio.sleep(backoff)
                continue

            except Exception as e:
                if self._sdk_debug:
                    print(f"[llm][err] {type(e).__name__}: {e}")
                last_exc = e
                if i < self._max_retries:
                    backoff = min(self._backoff_cap, self._backoff_base * (2**i))
                    backoff = backoff * (0.75 + 0.5 * random.random())
                    await asyncio.sleep(backoff)
                    continue
                break

        raise (last_exc or RuntimeError("LLM call failed"))

    # ----------------- 外部接口 -----------------
    def build_index(self, data_dir: str, work_dir: str) -> None:
        """
        构建/更新索引：
        - 递归扫描 data_dir 的 *.pdf / *.txt / *.md / *.docx
        - 分页/分块 -> embeddings -> 写入 Milvus-Lite
        """
        return self._run(self._build_index_async(data_dir, work_dir))

    def query(self, question: str) -> Dict[str, Any]:
        """
        检索增强回答：
        - embed 问题 -> Milvus TopK -> 拼接上下文 -> LLM 作答
        - 返回：answer/used_tokens/retrieved/contexts（供 RAGAS）
        """
        return self._run(self._query_async(question))

    # ----------------- 内部异步实现 -----------------
    async def _build_index_async(self, data_dir: str, work_dir: str) -> None:
        self._stage = "index"
        try:
            if not os.path.isdir(data_dir):
                raise RuntimeError(f"数据目录不存在: {data_dir}")

            # Milvus collection
            self._ensure_collection()
            m = self._milvus()

            # 扫描文件
            patterns = ["**/*.pdf", "**/*.txt", "**/*.md", "**/*.docx"]
            files: List[str] = []
            for pat in patterns:
                files.extend(glob.glob(os.path.join(data_dir, pat), recursive=True))
            files = sorted(set(files))

            if not files:
                print(f"[Engine] No files under {data_dir}, empty index")

            # 分批处理：页/文本 -> 切分 -> 嵌入 -> upsert
            batch_vectors: List[List[float]] = []
            batch_rows: List[Dict[str, Any]] = []
            total_chunks = 0

            for fp in files:
                # 读取为多条 (text, meta)
                try:
                    entries = _read_file_text(fp)
                except Exception as e:
                    print(f"[Engine] 读取失败，跳过: {fp} : {e}")
                    continue

                for text, meta in entries:
                    # 分块（token 滑窗）
                    chunks = _sliding_token_chunks(
                        text, self.tokenizer, self.chunk_tokens, self.chunk_overlap
                    )
                    # 针对超长 varchar，存储前可裁剪字符数（不影响用于上下文拼接时的原文，拼接再按 token 裁）
                    for idx, ch in enumerate(chunks):
                        chunk_text = ch[: self.text_max_chars]
                        row = {
                            "text": chunk_text,
                            "doc_id": meta.get("doc_id", ""),
                            "path": meta.get("path", ""),
                            "page": int(meta.get("page", -1)),
                            "chunk_id": int(idx),
                        }
                        batch_rows.append(row)
                        total_chunks += 1

                        # 批量阈值：按文本条数凑一批出 embedding
                        if len(batch_rows) >= self.embed_batch:
                            vecs = self._embed_batch([r["text"] for r in batch_rows])
                            for i in range(len(batch_rows)):
                                data = {
                                    "vector": vecs[i].tolist(),
                                    **batch_rows[i],
                                }
                                m.insert(self.collection, data=[data])
                            batch_rows.clear()

            # 处理尾巴
            if batch_rows:
                vecs = self._embed_batch([r["text"] for r in batch_rows])
                for i in range(len(batch_rows)):
                    data = {
                        "vector": vecs[i].tolist(),
                        **batch_rows[i],
                    }
                    m.insert(self.collection, data=[data])
                batch_rows.clear()

            print(f"[Engine] build_index done. files={len(files)} chunks={total_chunks} collection={self.collection}")

        finally:
            self._stage = "idle"

    async def _query_async(self, question: str) -> Dict[str, Any]:
        await self._ensure_llm_async()

        # 1) 嵌入问题
        q_vec = self._embed_batch([question])
        if q_vec.shape[0] == 0:
            raise RuntimeError("question embedding failed")
        qv = q_vec[0].tolist()

        # 2) Milvus TopK
        m = self._milvus()
        # output_fields 决定返回标量列
        res = m.search(
            self.collection,
            data=[qv],
            limit=self.top_k,
            output_fields=["text", "doc_id", "path", "page", "chunk_id"],
            search_params={"metric_type": self.metric_type},
        )
        # m.search 返回 list[Hits]，这里只取第一个 query 的结果
        hits = res[0] if res else []
        contexts: List[str] = []
        picked: List[Dict[str, Any]] = []

        for h in hits:
            score = float(getattr(h, "score", 0.0))
            row = h.fields  # 包含 text/doc_id/path/page/chunk_id
            if self.score_threshold > 0 and score < self.score_threshold:
                continue
            txt = str(row.get("text", "") or "")
            contexts.append(txt)
            picked.append({
                "text": txt,
                "doc_id": row.get("doc_id", ""),
                "path": row.get("path", ""),
                "page": int(row.get("page", -1)),
                "chunk_id": int(row.get("chunk_id", -1)),
                "score": score,
            })

        # 3) 组装上下文（拼接再按 token 裁到 context_max_tokens）
        #    为避免 prompt 爆表，这里按照加入顺序（已按向量相似度排序）
        ctx_joined = ""
        for c in contexts:
            if self.tokenizer.count(ctx_joined) >= self.context_max_tokens:
                break
            sep = "\n\n---\n\n" if ctx_joined else ""
            ctx_joined = ctx_joined + sep + c
        # token 裁剪
        ctx_joined = self.tokenizer.truncate(ctx_joined, self.context_max_tokens)

        # 4) 构造 Prompt & 调 LLM
        system_prompt = (
            "你是一个严格的企业检索问答助手。请仅依据“给定资料”，"
            "回答用户问题；若资料无法支持答案，请明确说明“资料不足，无法回答”。"
        )
        user_prompt = (
            f"【给定资料】\n{ctx_joined}\n\n"
            f"【问题】\n{question}\n\n"
            f"请基于给定资料作答，避免主观臆断；如无法从资料中得到，请说“资料不足，无法回答”。"
        )

        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ]

        answer = await self._llm_call(messages, max_tokens=512, temperature=0.2)

        return {
            "answer": answer,
            "used_tokens": int(self._last_usage.get("total_tokens", 0)),
            "retrieved": int(len(picked)),
            "contexts": contexts,        # 供 RAGAS 使用
            # 可选扩展供你调试：sources（不影响 evaluator）
            # "sources": picked,
        }
[root@node-0 naive-rag-sjs-algo-app]#

